import os
import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader #for PDFs
from docx import Document #for Word documents
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_groq import ChatGroq


#-----Text Extraction Functions------#
def extract_from_PDF(pdf): #function to extract text from PDF's
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()

    return text

def extract_from_Word(file): #function to extract text from Word Documents
    word_file = Document(file)
    text = "\n".join([paragraph.text for paragraph in word_file.paragraphs])
    
    return text

def extract_from_txt(file): #function to read from plain text files (notepad etc.)
    return file.read().decode('utf-8')


#------Dynamic Chunk Size Calculation/Text Splitting------#
def make_chunks(text):
    #determining ideal chunk size based on the input document. Base value established between 500-1500
    doc_length = len(text)
    base_size = min(max(int(doc_length * 0.01), 500), 1500) #500-1500 dependent on document length
    #assumption here is that longer documents (like research papers should have larger context windows compared to 
    #smaller documents like a conversation)

    #scaling chunk size according to word density
    words = text.split()
    avg_word_length = 0
    if words:
        avg_word_length = ((sum(len(w) for w in words)) / len(words)) #avg word length calculation
    else:
        avg_word_length = 5.0

    density_factor = max(1.0, avg_word_length / 5.0) #scaling chunk size according to word density of document
    chunk_size = int(base_size / density_factor)
    chunk_overlap = min(max(100, int(0.2 * chunk_size)), chunk_size // 2) 

    splitter = RecursiveCharacterTextSplitter( 
        chunk_size = chunk_size,
        chunk_overlap = chunk_overlap,
        separators = ["\n\n", "\n", " ", ""]
    )

    return splitter.split_text(text)

def main():
    load_dotenv()
    #webpage design
    st.set_page_config(page_title = "iDoc", layout = "wide")
    st.header("iDoc - Powered by Llama3 70B")

    #initialise DB
    embeddings = HuggingFaceBgeEmbeddings(model_name = "BAAI/bge-small-en-v1.5",
                                        model_kwargs = {"device": "cpu"},
                                        encode_kwargs = {"normalize_embeddings": True, "binary": True})
    
    #load existing DB or initliase a new one
    knowledge_base = Chroma(
        persist_directory = './chroma_db',
        embedding_function = embeddings,
        collection_metadata = {"hhssw:space": "hamming"}
    )

    with st.sidebar:
        st.subheader("Documents List")

        #load currently added documents (if any) from the DB
        collection = knowledge_base.get()
        existing_docs = collection['metadatas']
        doc_names = list(set(doc.get('source', 'Unknown') for doc in existing_docs)) if existing_docs else [] #extract document names

        if doc_names:
            selected_docs = st.multiselect(
                "Select document(s) to remove : ", 
                doc_names,
            )

            #remove multiple selected documents
            if st.button("Remove Selected Document(s)"):
                ids_to_remove = [collection['ids'][i] for i, doc in enumerate(existing_docs) if doc.get('source') in selected_docs]
                knowledge_base.delete(ids_to_remove)
                knowledge_base.persist()
                st.success(f"Removed {len(selected_docs)} document(s)!")
                st.rerun()
        else:
            st.info("No documents present in the database!")
        

    #upload files
    files = st.file_uploader(label = "Upload Documents:", type = ["pdf", "docx", "txt"], label_visibility = "hidden", accept_multiple_files = True)
    #check if valid file
    if files:
        #determine document type and then extract text accordingly
        valid_files = 0
        all_chunks = []
        for file in files:

            existing_sources = [doc.get('source') for doc in knowledge_base.get()['metadatas']]
            if file.name in existing_sources:
                st.warning(f"Skipped duplicate: {file.name}")
                continue

            text = ""
            if file.type == "application/pdf": #for PDF
                text = extract_from_PDF(file)
                # st.write(text)

            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": #for Word files
                text = extract_from_Word(file)
                # st.write(text)
            
            elif file.type == "text/plain": #for notepad etc.
                text = extract_from_txt(file)
                # st.write(text)

            else:
                st.error("Invalid file type entered!")
                continue

            #split extracted text into chunks
            chunks = make_chunks(text)
            if chunks:
                all_chunks.extend(chunks)
                valid_files += 1
                knowledge_base.add_texts(chunks, metadatas = [{"source" : file.name}] * len(chunks))
                knowledge_base.persist()
                st.success(f"Document of length {len(chunks)} chunks added!")

        
    if len(knowledge_base.get()['metadatas']) > 0: #if any chunks were made/files were processed
        
        query = st.text_input("Ask a question about your documents:") 
        
        llm = ChatGroq(model_name = "llama3-70b-8192", #set up LLM 
                api_key = os.getenv("GROQ_API_KEY"),
                temperature = 0.4) #modify temperature according to requirements. Lower means more accurate/less creative and vice versa.
        
        if query:
            with st.spinner("Searching your documents 🚀"):

                qa_chain = RetrievalQA.from_chain_type( #load LLM's response based on provided context (if any)
                    llm = llm,
                    chain_type = "stuff",
                    retriever = knowledge_base.as_retriever()
                )

                response = qa_chain.run(query)
                st.write(response)  #display response 
    else:
        st.error("Please upload documents to query!")

if __name__ == '__main__':
    main()

